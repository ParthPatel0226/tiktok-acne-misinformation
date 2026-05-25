"""
AAD + NLM Acne Guidelines Scraper v2
=====================================
Scrapes authoritative dermatology sources for the RAG pipeline ground truth.

Sources:
  - American Academy of Dermatology: https://www.aad.org/public/diseases/acne
  - National Library of Medicine / MedlinePlus: https://medlineplus.gov/acne.html

Output: acne_guidelines_2.docx (all scraped content combined)

Usage:
    pip install requests beautifulsoup4 python-docx
    python acne_scraper_v2.py --output acne_guidelines_2.docx --depth unlimited
    python acne_scraper_v2.py --output acne_guidelines_2.docx --depth 3
"""

import argparse
import time
import re
import sys
from collections import deque
from urllib.parse import urljoin, urlparse

try:
    import requests
    from bs4 import BeautifulSoup
except ImportError:
    print("ERROR: Install dependencies: pip install requests beautifulsoup4")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt
except ImportError:
    print("WARNING: python-docx not installed. Output will be .txt instead.")
    Document = None


HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
}

AAD_BASE = "https://www.aad.org/public/diseases/acne"
NLM_BASE = "https://medlineplus.gov/acne.html"

AAD_SEED_URLS = [
    "https://www.aad.org/public/diseases/acne",
    "https://www.aad.org/public/diseases/acne/causes",
    "https://www.aad.org/public/diseases/acne/treatment",
    "https://www.aad.org/public/diseases/acne/really-acne",
    "https://www.aad.org/public/diseases/acne/skin-care",
    "https://www.aad.org/public/diseases/acne/diy",
]

NLM_SEED_URLS = [
    "https://medlineplus.gov/acne.html",
    "https://medlineplus.gov/ency/article/000873.htm",
]


def fetch_page(url):
    """Fetch a page and return BeautifulSoup object."""
    try:
        resp = requests.get(url, headers=HEADERS, timeout=15)
        resp.raise_for_status()
        return BeautifulSoup(resp.text, "html.parser")
    except Exception as e:
        print(f"  Failed to fetch {url}: {e}")
        return None


def extract_text(soup):
    """Extract main content text from a page."""
    # Remove scripts, styles, nav
    for tag in soup.find_all(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    
    # Try main content area
    main = soup.find("main") or soup.find("article") or soup.find("div", class_="content")
    if main:
        text = main.get_text(separator="\n", strip=True)
    else:
        text = soup.get_text(separator="\n", strip=True)
    
    # Clean up
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    return "\n".join(lines)


def find_links(soup, base_url, domain_filter):
    """Find all links on a page within the same domain."""
    links = set()
    for a in soup.find_all("a", href=True):
        href = urljoin(base_url, a["href"])
        parsed = urlparse(href)
        clean = f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
        if domain_filter in clean and clean != base_url:
            links.add(clean)
    return links


def crawl_site(seed_urls, domain_filter, max_depth=None):
    """BFS crawl a site from seed URLs."""
    visited = set()
    pages = []
    queue = deque([(url, 0) for url in seed_urls])

    while queue:
        url, depth = queue.popleft()
        if url in visited:
            continue
        if max_depth is not None and depth > max_depth:
            continue
        
        visited.add(url)
        print(f"  [{len(visited)}] Crawling: {url}")
        
        soup = fetch_page(url)
        if not soup:
            continue
        
        title = soup.title.string.strip() if soup.title else url
        text = extract_text(soup)
        
        if len(text) > 100:  # Skip near-empty pages
            pages.append({"url": url, "title": title, "text": text, "depth": depth})
        
        # Find more links
        new_links = find_links(soup, url, domain_filter)
        for link in new_links:
            if link not in visited:
                queue.append((link, depth + 1))
        
        time.sleep(1)  # Polite delay

    return pages


def save_docx(pages, output_path):
    """Save all scraped content to a Word document."""
    doc = Document()
    doc.add_heading("Acne Clinical Guidelines — AAD + NLM", level=0)
    doc.add_paragraph(f"Scraped {len(pages)} pages for RAG pipeline ground truth.")
    doc.add_paragraph("")

    for page in pages:
        doc.add_heading(page["title"], level=1)
        doc.add_paragraph(f"Source: {page['url']}", style="Intense Quote")
        doc.add_paragraph(page["text"])
        doc.add_page_break()

    doc.save(output_path)
    print(f"\nSaved: {output_path} ({len(pages)} pages)")


def save_txt(pages, output_path):
    """Fallback: save as plain text."""
    with open(output_path, "w", encoding="utf-8") as f:
        for page in pages:
            f.write(f"{'='*60}\n")
            f.write(f"TITLE: {page['title']}\n")
            f.write(f"URL:   {page['url']}\n")
            f.write(f"{'='*60}\n\n")
            f.write(page["text"])
            f.write("\n\n")
    print(f"\nSaved: {output_path} ({len(pages)} pages)")


def main():
    parser = argparse.ArgumentParser(description="Scrape AAD + NLM acne guidelines")
    parser.add_argument("--output", default="acne_guidelines_2.docx", help="Output file")
    parser.add_argument("--depth", default="3", help="Max crawl depth (number or 'unlimited')")
    args = parser.parse_args()

    max_depth = None if args.depth == "unlimited" else int(args.depth)

    print("=" * 55)
    print("AAD + NLM Acne Guidelines Scraper v2")
    print("=" * 55)

    # Crawl AAD
    print(f"\n--- Crawling AAD (depth={args.depth}) ---")
    aad_pages = crawl_site(AAD_SEED_URLS, "aad.org/public/diseases/acne", max_depth)
    print(f"  AAD: {len(aad_pages)} pages")

    # Crawl NLM
    print(f"\n--- Crawling NLM/MedlinePlus ---")
    nlm_pages = crawl_site(NLM_SEED_URLS, "medlineplus.gov", max_depth=2)
    print(f"  NLM: {len(nlm_pages)} pages")

    all_pages = aad_pages + nlm_pages

    # Save
    if Document and args.output.endswith(".docx"):
        save_docx(all_pages, args.output)
    else:
        txt_path = args.output.replace(".docx", ".txt")
        save_txt(all_pages, txt_path)

    print(f"\nTotal pages scraped: {len(all_pages)}")
    print("Use this file as ground truth in your RAG pipeline.")


if __name__ == "__main__":
    main()
